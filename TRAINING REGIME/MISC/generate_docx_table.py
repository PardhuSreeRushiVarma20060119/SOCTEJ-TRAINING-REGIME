import re
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_curriculum(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    parsed_days = []
    seen_ids = set()
    current_day = None
    
    # Loose regex
    day_pattern = re.compile(r'\*\*Day (\d+(?:-\d+)?):? (.*?)\*\*')
    
    for line in lines:
        line = line.strip()
        if not line: continue
        
        match = day_pattern.search(line)
        if match:
            day_id = match.group(1)
            title = match.group(2).strip(': ')
            
            # Normalizing ID to handle duplicates/overrides
            if day_id in seen_ids:
                if current_day and day_id in current_day["Day"]:
                    if title and title not in current_day["Objective"]:
                        current_day["Tasks"].append(f"Focus: {title}")
                continue
            
            if current_day:
                parsed_days.append(current_day)
            
            seen_ids.add(day_id)
            current_day = {
                "Day": f"Day {day_id}",
                "Objective": title or "Tactical Operations",
                "Tasks": []
            }
            continue
        
        if current_day:
            if "**Objective:**" in line:
                obj = line.replace("**Objective:**", "").strip()
                if obj: current_day["Objective"] = obj
            elif line.startswith("- [ ]") or line.startswith("- "):
                current_day["Tasks"].append(line.replace("- [ ]", "").replace("- ", "").strip())
            elif ":" in line and any(c.isdigit() for c in line.split(':')[0]):
                current_day["Tasks"].append(line)

    if current_day:
        parsed_days.append(current_day)
    
    # Sort correctly
    def get_first_num(day_str):
        m = re.search(r'(\d+)', day_str)
        return int(m.group(1)) if m else 0
    
    parsed_days.sort(key=lambda x: get_first_num(x["Day"]))
    
    return parsed_days

def create_docx(data, output_path):
    doc = Document()
    title = doc.add_heading('112-DAY SOC ANALYST ELITE TRAINING REGIME', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Master Task Table - Complete 112-Day Curriculum')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Day'
    hdr_cells[1].text = 'Objective'
    hdr_cells[2].text = 'Tactical Tasks & Operations'
    
    for day in data:
        row_cells = table.add_row().cells
        row_cells[0].text = day['Day']
        row_cells[1].text = day['Objective']
        row_cells[2].text = "\n".join(day['Tasks'][:15]) # Increased line limit
        
    doc.save(output_path)
    print(f"File saved to {output_path}")

if __name__ == "__main__":
    source = r"c:\Users\pardh\Downloads\TEJU INTENSIVE TRAINING\TRAINING REGIME\Docs\SOC_Training_Program.md"
    output = r"c:\Users\pardh\Downloads\TEJU INTENSIVE TRAINING\SOC_112_DAY_TRAINING_REGIME_TABLE.docx"
    
    if os.path.exists(source):
        curriculum_data = parse_curriculum(source)
        all_days = [d["Day"] for d in curriculum_data]
        print(f"Total Unique Day Entries: {len(all_days)}")
        print("List of Days Found:", ", ".join(all_days))
        create_docx(curriculum_data, output)
    else:
        print(f"Error: Source file {source} not found.")
