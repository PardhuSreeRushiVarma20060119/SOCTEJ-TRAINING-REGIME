import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def parse_markdown_file(file_path):
    print(f"Parsing: {os.path.basename(file_path)}")
    with open(file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    parsed_days = {} # Key: Day ID (str), Value: dict
    current_day_id = None
    
    # Regex to capture Day X or Day X-Y
    day_regex = re.compile(r'\*\*Day (\d+(?:-\d+)?):? (.*?)\*\*')
    
    for line in lines:
        line = line.strip()
        if not line: continue
        
        # Check for Day Header
        match = day_regex.search(line)
        if match:
            current_day_id = match.group(1)
            title = match.group(2).strip(': ')
            
            if current_day_id not in parsed_days:
                parsed_days[current_day_id] = {
                    "Objective": set(),
                    "Tasks": [],
                    "Mappings": set()
                }
            
            if title:
                parsed_days[current_day_id]["Objective"].add(title)
            continue
        
        if current_day_id:
            # Objective explicit line
            if "**Objective:**" in line:
                obj = line.replace("**Objective:**", "").strip()
                if obj: parsed_days[current_day_id]["Objective"].add(obj)
            
            # Framework Mappings
            elif "**ATT&CK Mapping:**" in line or "**ATT&CK:**" in line:
                mapping = re.sub(r'\*\*ATT&CK.*?\*\* ?:?', '', line).strip()
                parsed_days[current_day_id]["Mappings"].add(f"ATT&CK: {mapping}")
            elif "**D3FEND Mapping:**" in line or "**D3FEND:**" in line or "D3FEND Countermeasure" in line:
                # Remove links and formatting
                clean_line = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', line) 
                mapping = re.sub(r'\*\*D3FEND.*?\*\* ?:?', '', clean_line).strip()
                 # Clean up bolding
                mapping = mapping.replace('**', '')
                parsed_days[current_day_id]["Mappings"].add(f"D3FEND: {mapping}")
            
            # Tasks / Hourly Ops
            elif line.startswith("- ") or line.startswith("* ") or (":" in line and any(c.isdigit() for c in line.split(':')[0])):
                # Clean task
                task = line.replace("- [ ]", "").replace("- ", "").replace("* ", "").strip()
                # If it's a sub-bullet (e.g., hourly op), keep indentation or formatting?
                # For table, flat list is usually better, or simple prefix
                if not task.startswith("**Day"): # Avoid capturing headers
                    parsed_days[current_day_id]["Tasks"].append(task)

    return parsed_days

def aggregate_data(file_paths):
    master_days = {} # Key: numeric sortable, Value: consolidated dict
    
    for fp in file_paths:
        if os.path.exists(fp):
            file_data = parse_markdown_file(fp)
            for day_id, content in file_data.items():
                if day_id not in master_days:
                    master_days[day_id] = {
                        "Objective": content["Objective"],
                        "Tasks": set(content["Tasks"]), # Use set for dedup
                        "Mappings": content["Mappings"]
                    }
                else:
                    master_days[day_id]["Objective"].update(content["Objective"])
                    master_days[day_id]["Tasks"].update(content["Tasks"])
                    master_days[day_id]["Mappings"].update(content["Mappings"])
        else:
            print(f"Warning: File not found: {fp}")

    # Convert to list and sort
    sorted_days = []
    
    def get_sort_key(did):
        # Handle ranges like "106-110" -> 106
        m = re.match(r'^(\d+)', did)
        return int(m.group(1)) if m else 999
        
    for day_id in sorted(master_days.keys(), key=get_sort_key):
        # Pick longest objective
        objs = list(master_days[day_id]["Objective"])
        best_obj = max(objs, key=len) if objs else "Standard Operations"
        
        # Sort tasks to keep some order (hourly ops first usually starts with digit)
        tasks = list(master_days[day_id]["Tasks"])
        tasks.sort() # Alphabetical sort puts timestamps 09:00 at top usually
        
        sorted_days.append({
            "Day": f"Day {day_id}",
            "Objective": best_obj,
            "Tasks": "\n".join(tasks),
            "Mappings": "\n".join(sorted(list(master_days[day_id]["Mappings"])))
        })
        
    return sorted_days

def create_docx(data, output_path):
    doc = Document()
    
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    
    # Title
    head = doc.add_heading('112-DAY SOC ANALYST ELITE TRAINING REGIME', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph('Consolidated Master Table (Multi-Source Aggregation)')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph('Sources: Standard Curriculum, MITRE ATT&CK/D3FEND Guides, Elite Corpus, & Master Regime.')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Day'
    hdr[1].text = 'Objective / Focus'
    hdr[2].text = 'Tactical Tasks & Hourly Ops'
    hdr[3].text = 'Framework Mappings'
    
    # Set widths (approximate)
    for cell in hdr:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), "D9D9D9") # Grey header
        tcPr.append(shd)
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell.text)
        run.bold = True

    for day in data:
        row = table.add_row().cells
        row[0].text = day['Day']
        row[1].text = day['Objective']
        row[2].text = day['Tasks']
        row[3].text = day['Mappings']
        
    doc.save(output_path)
    print(f"File saved to {output_path}")

if __name__ == "__main__":
    base_dir = r"c:\Users\pardh\Downloads\TEJU INTENSIVE TRAINING"
    
    sources = [
        os.path.join(base_dir, r"TRAINING REGIME\Docs\SOC_Training_Program.md"),
        os.path.join(base_dir, r"TRAINING REGIME\Docs (MITRE)\MitreATT&CK\SOC_Training_Program.md"),
        os.path.join(base_dir, r"ULTIMATE_SOC_TRAINING_REGIME_MASTER.md"),
        os.path.join(base_dir, r"SOC_ELITE_112_DAY_TRAINING_CORPUS.md")
    ]
    
    output_file = os.path.join(base_dir, "SOC_112_DAY_MASTER_CONSOLIDATED_TABLE.docx")
    
    print("Starting Aggregation...")
    data = aggregate_data(sources)
    print(f"Total Unique Days Aggregated: {len(data)}")
    create_docx(data, output_file)
